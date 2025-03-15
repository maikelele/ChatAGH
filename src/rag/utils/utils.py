import os
import json

import time
import functools

import pypdf
import docx
import glob
from dotenv import load_dotenv
from langchain.schema import Document


def load_data(path: str, max_docs: int = None):
    """
    Load documents from a file or directory path.

    Args:
        path: Path to a file or directory
        max_docs: Maximum number of documents to load (None for all)

    Returns:
        List of Document objects
    """
    documents = []

    if os.path.isdir(path):
        all_files = []
        for ext in ["*.json", "*.md", "*.txt", "*.pdf", "*.docx"]:
            all_files.extend(glob.glob(os.path.join(path, ext)))

        for file_path in all_files:
            documents.extend(load_file(file_path))
    else:
        documents = load_file(path)

    if max_docs is not None:
        documents = documents[:max_docs]

    return documents

def load_file(file_path: str):
    """
    Load a single file based on its extension

    Args:
        file_path: Path to the file

    Returns:
        List of Document objects
    """
    docs = []

    if file_path.endswith(".json"):
        with open(file_path) as f:
            data = json.load(f)

        if isinstance(data, list):
            docs = [Document(page_content=d.get("content", ""),
                             metadata=d.get("metadata", {})) for d in data]
        else:
            docs = [Document(page_content=data.get("content", ""),
                             metadata=data.get("metadata", {}))]

    elif file_path.endswith(".md"):
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        docs = [Document(page_content=content, metadata={"source": file_path})]

    elif file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        docs = [Document(page_content=content, metadata={"source": file_path})]

    elif file_path.endswith(".pdf"):
        docs = []
        pdf = pypdf.PdfReader(file_path)
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={
                        "source": file_path,
                        "page": i + 1,
                        "total_pages": len(pdf.pages)
                    }
                ))

    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        docs = [Document(page_content=content, metadata={"source": file_path})]

    return docs


def load_env():
    env_path = os.path.join(os.getcwd(), 'config', '.env')
    if os.path.exists(env_path):
        load_dotenv(dotenv_path=env_path)
    else:
        load_dotenv()


def retry_on_exception(attempts=3, delay=1, backoff=10, exception=Exception):
    """
    A decorator to retry a function call if it raises a specified exception.
    """
    def decorator(func):
        @functools.wraps(func)
        def wrapper(*args, **kwargs):
            current_delay = delay
            for attempt in range(1, attempts + 1):
                try:
                    return func(*args, **kwargs)
                except exception as e:
                    if attempt == attempts:
                        raise
                    else:
                        print(f"Attempt {attempt} failed: {e}. Retrying in {current_delay} seconds...")
                        time.sleep(current_delay)
                        current_delay *= backoff

        return wrapper

    return decorator
